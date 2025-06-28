from docx import Document
import io
import json
from docx.shared import Pt


def revised_document(file_contents, sharepoint_list):
    # Read the file_contents contents
    document = Document(io.BytesIO(file_contents))
    print(document.core_properties.title)

    sharepoint_list = json.loads(sharepoint_list)
    document = replace_placeholders(document, sharepoint_list)
    document = create_new_doc(document)

    # Save the Document to a BytesIO object
    fake_file = io.BytesIO()
    document.save(fake_file)

    fake_file.seek(0)  # Go to the beginning of the BytesIO object

    return fake_file


# do the replacement of the placeholders
def replace_placeholders(document, sharepoint_list):
    for item in sharepoint_list:
        placeholder = item["placeholder"]
        newtext = item["newtext"]

        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                # Clear the existing paragraph content and add the title
                paragraph.clear()
                paragraph.add_run(placeholder)

                # Add a new line and then the newtext
                newpara = document.add_paragraph()
                newpara.add_run(newtext)
                newpara.style = "Normal"

                # move the newpara after the original paragraph
                newpara._element.getparent().remove(newpara._element)
                paragraph._element.addnext(newpara._element)

    return document


def create_new_doc(document):

    # Create a new document
    new_document = Document()
    new_document.styles["Normal"].paragraph_format.space_after = Pt(6)

    # Copy the base styles from the original document to the new document
    for style in document.styles:
        if style.name not in new_document.styles:
            new_document.styles.add_style(style.name, style.type)
        # copy the font and size from the original document to the new document for this style
        try:
            new_document.styles[style.name].font.name = document.styles[
                style.name
            ].font.name
            new_document.styles[style.name].font.size = document.styles[
                style.name
            ].font.size
        except AttributeError:
            pass

    # Loop through all the paragraphs in the document and add them to the new document
    for paragraph in document.paragraphs:
        # if the paragraph contains "\n- " or starts with a number in paragraph.text, then call function process_list
        if "\n- " in paragraph.text or "\n1. " in paragraph.text or "\n* " in paragraph.text:
            new_document = process_list(new_document, paragraph)
        else:
            new_paragraph = new_document.add_paragraph(paragraph.text)
            new_paragraph.style = paragraph.style

    return new_document


def process_list(new_document, paragraph):
    # Split the paragraph into a list of lines
    lines = paragraph.text.split("\n")

    # Loop through all the lines in the list
    for line in lines:
        # If the line starts with a bullet point, add it to the new document as a bullet point
        if line.startswith("- "):
            # first remove the bullet point
            line = line[2:]
            new_paragraph = new_document.add_paragraph(line)
            new_paragraph.style = "List Bullet"
        elif len(line) > 0 and line[0].isdigit():  # Add a colon after "elif"
            # first remove the number
            line = line[2:].strip()
            new_paragraph = new_document.add_paragraph(line)
            new_paragraph.style = "List Number"
        # If the line doesn't start with a bullet point, add it to the new document as a normal paragraph
        else:
            new_paragraph = new_document.add_paragraph(line)
            new_paragraph.style = "Normal"

    return new_document  # Return the new document
