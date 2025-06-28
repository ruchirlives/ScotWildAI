from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import re


def apply_styles(doc_path, output_path):
    # Load the document
    doc = Document(doc_path)

    # Create styles if they do not exist
    styles = doc.styles

    # Create Heading 1 style
    if "Heading 1" not in styles:
        heading1 = styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
        heading1.font.name = "Calibri"
        heading1.font.size = Pt(18)
        heading1.font.bold = True

    # Create Heading 2 style
    if "Heading 2" not in styles:
        heading2 = styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
        heading2.font.name = "Calibri"
        heading2.font.size = Pt(14)
        heading2.font.bold = True

    # Create List Bullet style
    if "List Bullet" not in styles:
        list_bullet = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        list_bullet.font.name = "Calibri"
        list_bullet.font.size = Pt(12)

    # Create List Number style
    if "List Number" not in styles:
        list_number = styles.add_style("List Number", WD_STYLE_TYPE.PARAGRAPH)
        list_number.font.name = "Calibri"
        list_number.font.size = Pt(12)

    # Modify the Normal style
    normal_style = styles["Normal"]
    font = normal_style.font
    font.name = "Calibri"
    font.size = Pt(12)

    # Justify text and set spacing for Normal style
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    # Process paragraphs
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            continue
        elif is_list(paragraph):
            continue
        else:
            apply_normal_style(paragraph)

    # Process tables
    for table in doc.tables:
        handle_table(table)

    # Process images (inline shapes)
    for shape in doc.inline_shapes:
        handle_image(shape)

    # Save the modified document
    doc.save(output_path)


def apply_normal_style(paragraph):
    paragraph.style = "Normal"


def handle_table(table):
    # Process each cell in the table if needed
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if is_heading(paragraph):
                    continue
                elif is_list(paragraph):
                    continue
                else:
                    apply_normal_style(paragraph)


def handle_image(image):
    # No specific processing for images, just pass through
    pass


# Define heading styles based on some heuristic
def is_heading(paragraph):
    text = paragraph.text.strip()
    # Simple heuristic: All caps means Heading 1, otherwise Heading 2 if it ends with ':' or is a single line without a full stop
    if text.isupper():
        paragraph.style = "Heading 1"
    # check if the paragraph is bold or ends with a colon
    elif text.endswith(":") or (
        (len(text.splitlines()) == 1) and paragraph.runs[0].bold and not text.endswith(".")
    ):
        paragraph.style = "Heading 2"
    else:
        return False
    return True


# Define list item based on starting character
def is_list(paragraph):
    text = paragraph.text.strip()
    if re.match(r"^\d+\.\s", text):
        paragraph.style = "List Number"
        return True
    elif re.match(r"^[\u2022-\u2023]\s", text):
        paragraph.style = "List Bullet"
        return True
    return False


# Example usage
apply_styles("input.docx", "output.docx")
